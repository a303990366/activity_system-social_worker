import docx
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import os
import time


#if __name__ == '__main__':
def Word_print_two():
    yourPath = 'C:/Users/ACER/Desktop/圖片文件生成/資料'
    title_doc=input('標題名稱:')
    second_title=input('副標題:')

    allFileList = os.listdir(yourPath)
    for items in allFileList:
        doc = docx.Document()

        section = doc.sections[0]
        section.left_margin=Cm(1.27)

        section.right_margin=Cm(1.27)

        section.top_margin=Cm(1.27)

        section.bottom_margin=Cm(1.27)


        p = doc.add_paragraph()
        run = p.add_run(title_doc)
        run.font.size = Pt(20)
        p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        run = p.add_run(second_title)
        run.font.size = Pt(15)
        p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER


        table = doc.add_table(rows=2, cols=2)
        
        second_path=os.path.join(yourPath,items)
        pic_list=os.listdir(second_path)
        for j in range(0,2):
            for i in range(0,2):
                item1=random.choice(pic_list)
                pic_list.remove(item1)
                new_path=os.path.join(second_path,item1)
                cell=table.cell(i,j)
                c_p1 =cell.paragraphs[0]
                c_run1=c_p1.add_run()
                c_run1.add_picture(new_path,width=Cm(9.24),height=Cm(10.32))
        doc.save('C:/Users/ACER/Desktop/圖片文件生成/%s.docx' % items)
        print('已完成製作{0}文件.....'.format(items))
